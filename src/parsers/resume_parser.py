"""Resume parser for PDF and DOCX formats."""

import re
from pathlib import Path
from typing import List, Optional, Dict, Tuple
import pdfplumber
import docx2txt
from docx import Document

from src.models import (
    Resume,
    ContactInfo,
    Experience,
    Education,
    Project,
    Certification
)


class ResumeParser:
    """Parser for resume documents (PDF and DOCX)."""

    # Common section headers
    SECTION_HEADERS = {
        'contact': ['contact', 'personal information', 'contact information'],
        'summary': ['summary', 'professional summary', 'objective', 'profile', 'about'],
        'experience': ['experience', 'work experience', 'professional experience', 'employment', 'work history'],
        'education': ['education', 'academic background', 'qualifications'],
        'skills': ['skills', 'technical skills', 'core competencies', 'expertise', 'technologies','core technical skills'],
        'genai_skills': ['gen ai skill set', 'genai skills', 'ai skills', 'ml skills', 'ai/ml skills', 'genai & machine learning skills'],
        'projects': ['projects', 'personal projects', 'key projects', 'project highlights'],
        'certifications': ['certifications', 'certificates', 'licenses']
    }

    def __init__(self):
        """Initialize the resume parser."""
        pass

    def parse(self, file_path: str) -> Resume:
        """
        Parse a resume file (PDF or DOCX) into structured data.

        Args:
            file_path: Path to the resume file

        Returns:
            Resume object with structured data
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")

        # Extract text and hyperlinks based on file type
        hyperlinks = []
        if file_path.suffix.lower() == '.pdf':
            text = self._extract_text_from_pdf(str(file_path))
            print(f"[PARSER] Extracting hyperlinks from PDF...")
            hyperlinks = self._extract_hyperlinks_from_pdf(str(file_path))
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            text = self._extract_text_from_docx(str(file_path))
            print(f"[PARSER] Extracting hyperlinks from DOCX...")
            hyperlinks = self._extract_hyperlinks_from_docx(str(file_path))
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")

        print(f"[PARSER] Found {len(hyperlinks)} hyperlinks in document")

        # Parse the text into structured sections
        print(f"[PARSER] Identifying sections...")
        sections = self._identify_sections(text)
        print(f"[PARSER] Found sections: {list(sections.keys())}")

        # Extract structured data from each section
        print(f"[PARSER] Parsing contact info...")
        contact = self._parse_contact(sections.get('contact', '') + '\n' + text[:500], hyperlinks)

        print(f"[PARSER] Parsing summary...")
        summary = self._parse_summary(sections.get('summary', ''))

        print(f"[PARSER] Parsing experience...")
        experience = self._parse_experience(sections.get('experience', ''))

        print(f"[PARSER] Parsing education...")
        education = self._parse_education(sections.get('education', ''))

        print(f"[PARSER] Parsing skills...")
        skills = self._parse_skills(sections.get('skills', ''))

        print(f"[PARSER] Parsing GenAI skills...")
        genai_skills = self._parse_genai_skills(sections.get('genai_skills', ''))

        print(f"[PARSER] Parsing projects...")
        projects = self._parse_projects(sections.get('projects', ''))

        print(f"[PARSER] Parsing certifications...")
        certifications = self._parse_certifications(sections.get('certifications', ''))

        print(f"[PARSER] Creating Resume object...")
        resume = Resume(
            contact=contact,
            summary=summary,
            experience=experience,
            education=education,
            skills=skills,
            genai_skills=genai_skills,
            projects=projects,
            certifications=certifications,
            raw_text=text,
            source_file=str(file_path)
        )

        print(f"[PARSER] Resume parsing complete!")
        return resume

    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file with fallback."""
        text = ""

        # Try pdfplumber first (better formatting)
        try:
            print(f"[PDF PARSER] Trying pdfplumber on: {file_path}")
            with pdfplumber.open(file_path) as pdf:
                print(f"[PDF PARSER] PDF has {len(pdf.pages)} pages")
                for page_num, page in enumerate(pdf.pages, 1):
                    print(f"[PDF PARSER] Extracting page {page_num}/{len(pdf.pages)}")
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            print(f"[PDF PARSER] pdfplumber extracted {len(text)} characters")
        except Exception as e:
            print(f"[PDF PARSER] pdfplumber failed: {str(e)}")
            print(f"[PDF PARSER] Trying PyPDF2 as fallback...")

            # Fallback to PyPDF2
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    print(f"[PDF PARSER] PyPDF2 found {len(pdf_reader.pages)} pages")
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        print(f"[PDF PARSER] PyPDF2 extracting page {page_num}/{len(pdf_reader.pages)}")
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                print(f"[PDF PARSER] PyPDF2 extracted {len(text)} characters")
            except Exception as e2:
                print(f"[PDF PARSER ERROR] Both parsers failed. pdfplumber: {str(e)}, PyPDF2: {str(e2)}")
                raise ValueError(f"Error extracting text from PDF. Both pdfplumber and PyPDF2 failed: {e2}")

        if not text.strip():
            raise ValueError("PDF appears to be empty or contains only images")

        return text.strip()

    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            # Try using docx2txt first (simpler)
            text = docx2txt.process(file_path)
            if text and text.strip():
                return text.strip()

            # Fallback to python-docx
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text.strip()
        except Exception as e:
            raise ValueError(f"Error extracting text from DOCX: {e}")

    def _extract_hyperlinks_from_docx(self, file_path: str) -> List[Dict[str, str]]:
        """Extract all hyperlinks from DOCX file."""
        hyperlinks = []
        try:
            doc = Document(file_path)

            # Iterate through all paragraphs
            for paragraph in doc.paragraphs:
                # Check for hyperlinks in paragraph
                for run in paragraph.runs:
                    # Check if run has hyperlink
                    if run._element.tag.endswith('hyperlink'):
                        # Get hyperlink URL from relationship
                        hyperlink_id = run._element.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                        if hyperlink_id:
                            url = doc.part.rels[hyperlink_id].target_ref
                            hyperlinks.append({
                                'text': run.text,
                                'url': url
                            })

                # Also check hyperlinks at paragraph level
                for hyperlink in paragraph._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink'):
                    hyperlink_id = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    if hyperlink_id and hyperlink_id in doc.part.rels:
                        url = doc.part.rels[hyperlink_id].target_ref
                        text = ''.join(node.text for node in hyperlink.iter() if node.text)
                        hyperlinks.append({
                            'text': text,
                            'url': url
                        })

            print(f"[DOCX HYPERLINKS] Extracted {len(hyperlinks)} hyperlinks")
            for link in hyperlinks:
                print(f"  - {link['text']}: {link['url']}")

        except Exception as e:
            print(f"[DOCX HYPERLINKS] Error extracting hyperlinks: {e}")

        return hyperlinks

    def _extract_hyperlinks_from_pdf(self, file_path: str) -> List[Dict[str, str]]:
        """Extract all hyperlinks from PDF file."""
        hyperlinks = []
        try:
            import PyPDF2

            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                for page_num, page in enumerate(pdf_reader.pages):
                    # Extract annotations (links)
                    if '/Annots' in page:
                        annotations = page['/Annots']
                        for annotation in annotations:
                            obj = annotation.get_object()
                            if obj.get('/Subtype') == '/Link':
                                # Check for URI (external link)
                                if '/A' in obj and '/URI' in obj['/A']:
                                    url = obj['/A']['/URI']
                                    # Try to get the link text from Rect coordinates
                                    text = url  # Default to URL if no text found
                                    hyperlinks.append({
                                        'text': text,
                                        'url': url
                                    })

            print(f"[PDF HYPERLINKS] Extracted {len(hyperlinks)} hyperlinks")
            for link in hyperlinks:
                print(f"  - {link['text']}: {link['url']}")

        except Exception as e:
            print(f"[PDF HYPERLINKS] Error extracting hyperlinks: {e}")

        return hyperlinks

    def _identify_sections(self, text: str) -> Dict[str, str]:
        """
        Identify and extract different sections from resume text.

        Returns:
            Dictionary mapping section names to their content
        """
        sections = {}
        lines = text.split('\n')
        current_section = None
        section_content = []

        for line in lines:
            line_stripped = line.strip()
            line_lower = line_stripped.lower()

            # Check if this line is a section header
            detected_section = None

            # Don't detect section headers in bullet points, very long lines, or single-word continuations
            if (not line_stripped.startswith(('•', '-', '*', '▪', '●', '○')) and
                len(line_lower) < 50 and
                not line_stripped.endswith('.')):  # Exclude sentence continuations like "projects."

                for section_name, headers in self.SECTION_HEADERS.items():
                    for header in headers:
                        # More strict matching: header should be the main content of the line
                        # Either exact match or header at start followed by : or similar
                        # Exclude lines that are just the header word with a period
                        if line_lower == header + '.':
                            # Skip "projects." or "education." etc. - these are sentence continuations
                            continue

                        if (line_lower == header or
                            line_lower.startswith(header + ':') or
                            line_lower == header.upper() or  # All caps header
                            line_lower == header.title()):  # Title case header
                            detected_section = section_name
                            break
                    if detected_section:
                        break

            if detected_section:
                # Save previous section
                if current_section and section_content:
                    sections[current_section] = '\n'.join(section_content)

                # Start new section
                current_section = detected_section
                section_content = []
            elif current_section:
                # Add to current section
                section_content.append(line)

        # Save last section
        if current_section and section_content:
            sections[current_section] = '\n'.join(section_content)

        return sections

    def _parse_contact(self, text: str, hyperlinks: List[Dict[str, str]] = None) -> ContactInfo:
        """Parse contact information from text and extracted hyperlinks."""
        contact = ContactInfo()
        if hyperlinks is None:
            hyperlinks = []

        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact.email = email_match.group(0)

        # Extract phone (more specific pattern)
        phone_pattern = r'[\+\(]?[1-9][0-9 .\-\(\)]{8,}[0-9]'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            phone = phone_match.group(0).strip()
            # Validate it's a phone number (not a random number sequence)
            if len(re.sub(r'[^\d]', '', phone)) >= 10:
                contact.phone = phone

        # First, try to extract LinkedIn and GitHub from hyperlinks (more accurate)
        for link in hyperlinks:
            url = link['url'].lower()
            if 'linkedin.com/in/' in url and not contact.linkedin:
                # Extract full URL
                contact.linkedin = link['url']
                print(f"[CONTACT] LinkedIn from hyperlink: {contact.linkedin}")
            elif 'github.com/' in url and not contact.github:
                # Extract full URL
                contact.github = link['url']
                print(f"[CONTACT] GitHub from hyperlink: {contact.github}")

        # Fallback: Extract LinkedIn from text (full URL or just username)
        if not contact.linkedin:
            linkedin_pattern = r'(?:https?://)?(?:www\.)?linkedin\.com/in/([\w-]+)'
            linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
            if linkedin_match:
                contact.linkedin = f"https://linkedin.com/in/{linkedin_match.group(1)}"
            else:
                # Look for standalone "LinkedIn" followed by a colon and a link or username
                linkedin_simple = re.search(r'LinkedIn\s*:\s*(https?://[^\s|]+|[\w-]{3,})', text, re.IGNORECASE)
                if linkedin_simple:
                    link = linkedin_simple.group(1).strip()
                    if link.startswith('http'):
                        contact.linkedin = link
                    elif link.lower() not in ['github', 'link', 'profile', 'summary', 'professional', 'email']:
                        contact.linkedin = f"https://linkedin.com/in/{link}"

        # Fallback: Extract GitHub from text (full URL or just username)
        if not contact.github:
            github_pattern = r'(?:https?://)?(?:www\.)?github\.com/([\w-]+)'
            github_match = re.search(github_pattern, text, re.IGNORECASE)
            if github_match:
                contact.github = f"https://github.com/{github_match.group(1)}"
            else:
                # Look for standalone "GitHub" followed by a colon and a link or username
                github_simple = re.search(r'GitHub\s*:\s*(https?://[^\s|]+|[\w-]{3,})', text, re.IGNORECASE)
                if github_simple:
                    link = github_simple.group(1).strip()
                    if link.startswith('http'):
                        contact.github = link
                    elif link.lower() not in ['linkedin', 'link', 'profile', 'summary', 'professional', 'email']:
                        contact.github = f"https://github.com/{link}"

        # Extract name (usually first line or before email)
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        if lines:
            # Name is likely the first substantial line (all caps often indicates name in resumes)
            for line in lines[:5]:
                # Skip lines with @ (emails), phone numbers, or common labels
                if any(x in line.lower() for x in ['@', 'phone', 'email', 'linkedin', 'github', 'summary', 'objective']):
                    continue
                # Look for all-caps name or title case name
                if (line.isupper() or line.istitle()) and 3 < len(line) < 50:
                    # Check it's not a section header
                    if not any(x in line.lower() for x in ['experience', 'education', 'skills', 'professional', 'summary']):
                        contact.full_name = line
                        break

        # Extract location (look for city, state patterns)
        # Be more careful to avoid matching technology names
        # Location typically appears near contact info, not in skills section
        contact_section = text[:500]  # First 500 chars likely contain contact info

        # Look for explicit location indicators
        location_patterns = [
            r'Location[\s:]+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?,\s*[A-Z]{2}(?:,\s*[A-Z]{2,})?)',  # "Location: City Name, ST"
            r'\|\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?,\s*[A-Z]{2})\s*\|',  # "| City Name, ST |" (pipe-separated)
            r'\|\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?,\s*[A-Z]{2})\s*$',  # "| City Name, ST" at end of line
            r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?,\s*[A-Z]{2}(?:,\s*(?:USA|United States)))\b'  # City Name, ST, USA
        ]

        for pattern in location_patterns:
            location_match = re.search(pattern, contact_section, re.MULTILINE)
            if location_match:
                contact.location = location_match.group(1).strip()
                break

        return contact

    def _parse_summary(self, text: str) -> Optional[str]:
        """Parse professional summary."""
        text = text.strip()
        return text if text else None

    def _parse_experience(self, text: str) -> List[Experience]:
        """Parse work experience entries."""
        experiences = []

        if not text.strip():
            return experiences

        lines = text.split('\n')
        current_exp = None

        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue

            # Pattern 1: "Job Title – Company, Location | Date – Date" (all on one line)
            job_match_inline = re.search(
                r'^([^–—\-|]+?)\s*[–—-]\s*([^,|]+(?:,\s*[^|]+)?)\s*\|\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4})\s*[–—-]\s*(Present|Current|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4})',
                line,
                re.IGNORECASE
            )

            # Pattern 2: "Company, Location | Date – Date" (job title is on previous line)
            job_match_multiline = re.search(
                r'^([^|]+?),\s*([^|]+?)\s*\|\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4})\s*[–—-]\s*(Present|Current|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4})',
                line,
                re.IGNORECASE
            )

            if job_match_inline:
                # Save previous experience
                if current_exp and (current_exp.company or current_exp.title):
                    experiences.append(current_exp)

                # Pattern 1: All on one line
                title = job_match_inline.group(1).strip()
                company_loc = job_match_inline.group(2).strip()
                start_date = job_match_inline.group(3).strip()
                end_date = job_match_inline.group(4).strip()

                # Split company and location
                company_parts = company_loc.split(',')
                company = company_parts[0].strip()
                location = ', '.join(company_parts[1:]).strip() if len(company_parts) > 1 else None

                current_exp = Experience(
                    title=title,
                    company=company,
                    location=location,
                    start_date=start_date,
                    end_date=end_date,
                    is_current=(end_date.lower() in ['present', 'current']),
                    bullets=[]
                )

            elif job_match_multiline:
                # Save previous experience
                if current_exp and (current_exp.company or current_exp.title):
                    experiences.append(current_exp)

                # Pattern 2: Company on current line, title on previous line
                # Look back for job title on previous non-empty line
                title = None
                for j in range(i-1, max(0, i-5), -1):
                    prev_line = lines[j].strip()
                    if prev_line and not prev_line.startswith(('•', '-', '*', '▪', '●', '○')):
                        # Check if it looks like a job title (not a bullet or date line)
                        if not re.search(r'\d{4}', prev_line) and '|' not in prev_line:
                            title = prev_line
                            break

                company = job_match_multiline.group(1).strip()
                location = job_match_multiline.group(2).strip()
                start_date = job_match_multiline.group(3).strip()
                end_date = job_match_multiline.group(4).strip()

                current_exp = Experience(
                    title=title if title else "Unknown",
                    company=company,
                    location=location,
                    start_date=start_date,
                    end_date=end_date,
                    is_current=(end_date.lower() in ['present', 'current']),
                    bullets=[]
                )

            else:
                # Pattern 3: Multi-line format without pipe separator
                # Line 1: Company, Location
                # Line 2: Job Title
                # Line 3: Date – Date
                # Check if current line is a date range (look ahead for title and company)
                date_match = re.search(
                    r'^((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4})\s*[–—-]\s*(Present|Current|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4})',
                    line,
                    re.IGNORECASE
                )

                if date_match and i >= 2:
                    # Look back for title and company, skipping empty lines
                    title_line = None
                    company_line = None

                    # Find the last 2 non-empty lines before current position
                    non_empty_lines = []
                    for j in range(i-1, max(-1, i-10), -1):
                        if j >= 0 and j < len(lines):
                            prev_line = lines[j].strip()
                            if prev_line and not prev_line.startswith(('•', '-', '*', '▪', '●', '○')):
                                non_empty_lines.append(prev_line)
                                if len(non_empty_lines) == 2:
                                    break

                    # Assign title and company from non-empty lines
                    if len(non_empty_lines) >= 2:
                        title_line = non_empty_lines[0]  # Most recent non-empty line (title)
                        company_line = non_empty_lines[1]  # Second most recent (company)

                    # Check if this looks like a valid experience entry
                    if (title_line and company_line and
                        not re.search(r'\d{4}', title_line)):  # Title shouldn't have year

                        # Save previous experience
                        if current_exp and (current_exp.company or current_exp.title):
                            experiences.append(current_exp)

                        # Parse company and location
                        company_parts = company_line.split(',')
                        company = company_parts[0].strip()
                        location = ', '.join(company_parts[1:]).strip() if len(company_parts) > 1 else None

                        start_date = date_match.group(1).strip()
                        end_date = date_match.group(2).strip()

                        current_exp = Experience(
                            title=title_line,
                            company=company,
                            location=location,
                            start_date=start_date,
                            end_date=end_date,
                            is_current=(end_date.lower() in ['present', 'current']),
                            bullets=[]
                        )
                        continue  # Skip to next line

            if current_exp:
                # Check if this is a bullet point
                if line.startswith(('•', '-', '*', '▪', '●', '○')):
                    # This is a new bullet point
                    cleaned = re.sub(r'^[•\-*▪●○]\s*', '', line)
                    if cleaned:
                        current_exp.bullets.append(cleaned)
                elif current_exp.bullets and line and not re.match(r'^[A-Z][a-z]+\s+[A-Z][a-z]+', line):
                    # This is a continuation of the previous bullet (not a new section header)
                    # Append to the last bullet
                    current_exp.bullets[-1] += ' ' + line
                elif not line.startswith(('•', '-', '*')) and '|' not in line and not re.search(r'\d{4}', line):
                    # No bullets yet, might be a description line - treat as a bullet
                    if len(line) > 20:  # Substantial content
                        current_exp.bullets.append(line)

        # Add last experience
        if current_exp and (current_exp.company or current_exp.title):
            experiences.append(current_exp)

        return experiences

    def _parse_education(self, text: str) -> List[Education]:
        """Parse education entries with timeout protection."""
        if not text.strip():
            return []

        # Use simple parser for now (detailed parser has regex issues causing hangs)
        print(f"[PARSER] Using simple education parser")
        return self._parse_education_simple(text)

    def _parse_education_simple(self, text: str) -> List[Education]:
        """Simple fallback education parser with basic extraction."""
        education_list = []
        lines = [line.strip() for line in text.split('\n') if line.strip()]

        i = 0
        while i < len(lines):
            line = lines[i].strip()

            # Look for university/college keywords
            if any(keyword in line.lower() for keyword in ['university', 'college', 'institute', 'school']):
                # Extract institution and location
                institution = line.split(',')[0].strip() if ',' in line else line
                location = line.split(',')[1].strip() if ',' in line and len(line.split(',')) > 1 else None

                # Look at next line for degree
                degree = None
                field_of_study = None
                graduation_date = None

                if i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    # Check if next line contains degree keywords
                    if any(deg in next_line.lower() for deg in ['master', 'bachelor', 'associate', 'phd', 'doctorate', 'm.s', 'b.s', 'mba']):
                        # Extract degree and field
                        if ' in ' in next_line.lower():
                            parts = next_line.split(' in ', 1)
                            degree = parts[0].strip()
                            field_of_study = parts[1].strip()
                        elif ' of ' in next_line.lower():
                            parts = next_line.split(' of ', 1)
                            degree = parts[0].strip()
                            field_of_study = parts[1].strip()
                        else:
                            degree = next_line
                        i += 1  # Skip the degree line

                        # Check next line for graduation date
                        if i + 1 < len(lines):
                            date_line = lines[i + 1].strip()
                            # Simple date extraction (look for 4-digit year)
                            import re
                            year_match = re.search(r'\b(19|20)\d{2}\b', date_line)
                            if year_match:
                                graduation_date = year_match.group(0)
                                i += 1  # Skip the date line

                edu = Education(
                    institution=institution,
                    degree=degree,
                    field_of_study=field_of_study,
                    graduation_date=graduation_date,
                    location=location,
                    gpa=None
                )
                education_list.append(edu)

            i += 1

        return education_list[:5]  # Limit to 5 entries

    def _parse_education_detailed(self, text: str) -> List[Education]:
        """Detailed education parser (original implementation)."""
        education_list = []

        if not text.strip():
            return education_list

        # Split by bullet points or multiple newlines
        lines = [line.strip() for line in text.split('\n') if line.strip()]

        # Track pending data for multi-line education entries
        pending_graduation_date = None
        pending_institution = None
        pending_location = None

        i = 0
        while i < len(lines):
            line = lines[i]
            # Remove bullet points
            line = re.sub(r'^[•\-*▪●○]\s*', '', line)

            # NEW PATTERN: Multi-line format
            # Line 1: "George Mason University, Fairfax, VA"
            # Line 2: "Master of Science in Information Technology"
            # Line 3: "August 2022 – May 2024"

            # Check if line contains institution (University, College, etc.) with location
            inst_loc_match = re.match(
                r'^((?:The\s+)?(?:.*?\s+)?(?:University|College|Institute|School)[^,]*),\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?,\s*[A-Z]{2})',
                line,
                re.IGNORECASE
            )

            if inst_loc_match and i + 1 < len(lines):
                pending_institution = inst_loc_match.group(1).strip()
                pending_location = inst_loc_match.group(2).strip()

                # Check next line for degree
                next_line = lines[i + 1].strip()
                degree_match = re.match(
                    r'^((?:Master|Bachelor|Associate|Doctor)(?:\s+of\s+(?:Science|Arts|Engineering|Business Administration))?|M\.S\.|B\.S\.|M\.A\.|B\.A\.|MBA|PhD|Ph\.D\.)\s+(?:in\s+)?(.+?)$',
                    next_line,
                    re.IGNORECASE
                )

                if degree_match:
                    degree = degree_match.group(1).strip()
                    field = degree_match.group(2).strip() if degree_match.group(2) else None

                    # Check line after for dates
                    grad_date = None
                    if i + 2 < len(lines):
                        date_line = lines[i + 2].strip()
                        # Match "Month YYYY – Month YYYY" or "Month YYYY - Month YYYY"
                        date_match = re.search(r'([A-Z][a-z]+\s+\d{4})\s*[-–—]\s*([A-Z][a-z]+\s+\d{4})', date_line)
                        if date_match:
                            grad_date = date_match.group(2)  # End date
                            i += 1  # Skip date line

                    edu = Education(
                        institution=pending_institution,
                        degree=degree,
                        field_of_study=field,
                        location=pending_location,
                        graduation_date=grad_date
                    )
                    education_list.append(edu)
                    pending_institution = None
                    pending_location = None
                    i += 2  # Skip institution and degree lines
                    continue

            # Check if this is a "Graduated: MM/YYYY" or "Graduation Date: MM/YYYY" line
            grad_date_match = re.search(r'(?:Graduated|Graduation Date):\s*(\d{2}/\d{4}|\d{4})', line, re.IGNORECASE)
            if grad_date_match:
                pending_graduation_date = grad_date_match.group(1)
                # If we already have an education entry, add the date to it
                if education_list:
                    education_list[-1].graduation_date = pending_graduation_date
                    pending_graduation_date = None
                i += 1
                continue

            # Pattern 0: "Degree in Field - Institution" (New pattern for this format)
            # Example: "Master of Science in Computer Science - University of Texas at Arlington"
            edu_match_dash = re.search(
                r'^([Mm]aster\s+of\s+[Ss]cience|[Bb]achelor\s+of\s+[Ss]cience|[Mm]aster\s+of\s+[Aa]rts|[Bb]achelor\s+of\s+[Aa]rts|M\.S\.|B\.S\.|M\.A\.|B\.A\.|MBA|PhD|Ph\.D\.)\s+(?:in|:)?\s+([^-]+?)\s*[-–—]\s*(.+)$',
                line,
                re.IGNORECASE
            )

            if edu_match_dash:
                degree = edu_match_dash.group(1).strip()
                field = edu_match_dash.group(2).strip()
                institution = edu_match_dash.group(3).strip()

                edu = Education(
                    institution=institution,
                    degree=degree,
                    field_of_study=field,
                    graduation_date=pending_graduation_date
                )
                education_list.append(edu)
                pending_graduation_date = None
                continue

            # Pattern 1: "Date Degree: Field, Institution - Location"
            # Example: "05/2024 Master of Science: Computer Science, The University of Texas at Arlington - Arlington, TX, USA"
            edu_match = re.search(
                r'(?:(\d{2}/\d{4})\s+)?([Mm]aster\s+of\s+[Ss]cience|[Bb]achelor\s+of\s+[Ss]cience|[Mm]aster\s+of\s+[Aa]rts|[Bb]achelor\s+of\s+[Aa]rts|M\.S\.|B\.S\.|M\.A\.|B\.A\.|MBA|PhD|Ph\.D\.)(?:\s+in)?\s*:?\s*([^,]+)?,?\s*(.+?)(?:\s*[-–—]\s*(.+))?$',
                line,
                re.IGNORECASE
            )

            if edu_match:
                date = edu_match.group(1)
                degree = edu_match.group(2).strip()
                field = edu_match.group(3).strip() if edu_match.group(3) else None
                institution_and_loc = edu_match.group(4).strip() if edu_match.group(4) else ""
                location = edu_match.group(5).strip() if edu_match.group(5) else None

                # Clean up institution (remove "in", ":", etc.)
                institution = re.sub(r'^[:\s,]+', '', institution_and_loc)

                # If field has ":" at the end, remove it
                if field:
                    field = field.rstrip(':,')

                edu = Education(
                    institution=institution,
                    degree=degree,
                    field_of_study=field,
                    location=location,
                    graduation_date=date
                )
                education_list.append(edu)
                i += 1
            else:
                # Pattern 2: Try simpler formats
                # Look for degree keywords
                if any(deg in line.lower() for deg in ['bachelor', 'master', 'phd', 'ph.d', 'b.s.', 'm.s.', 'b.a.', 'm.a.', 'degree', 'diploma']):
                    edu = Education(institution="")

                    # Extract graduation date (usually MM/YYYY or YYYY)
                    date_match = re.search(r'(\d{2}/\d{4}|\d{4})', line)
                    if date_match:
                        edu.graduation_date = date_match.group(1)
                        # Remove date from line for further parsing
                        line = line.replace(date_match.group(0), '').strip()

                    # Extract GPA
                    gpa_match = re.search(r'GPA:?\s*([\d.]+)', line, re.IGNORECASE)
                    if gpa_match:
                        edu.gpa = gpa_match.group(1)

                    # Extract degree (full form)
                    degree_match = re.search(
                        r'(Bachelor\s+of\s+Science|Master\s+of\s+Science|Bachelor\s+of\s+Arts|Master\s+of\s+Arts|PhD|Ph\.D\.|B\.S\.|M\.S\.|B\.A\.|M\.A\.)',
                        line,
                        re.IGNORECASE
                    )
                    if degree_match:
                        edu.degree = degree_match.group(0)

                    # Try to extract field of study (usually after "in" or after degree)
                    field_match = re.search(r'(?:in|:)\s+([A-Z][a-zA-Z\s]+?)(?:,|\s+-|\s+at|$)', line)
                    if field_match:
                        edu.field_of_study = field_match.group(1).strip()

                    # Try to extract institution (often contains "University" or "College" or "Institute")
                    inst_match = re.search(r'((?:The\s+)?(?:University|College|Institute|School)\s+[^,\-|]+)', line, re.IGNORECASE)
                    if inst_match:
                        edu.institution = inst_match.group(1).strip()
                    else:
                        # Use the whole line cleaned up
                        cleaned = re.sub(r'\d{2}/\d{4}|\d{4}|GPA:?\s*[\d.]+', '', line).strip()
                        if cleaned and len(cleaned) > 5:
                            edu.institution = cleaned

                    # Extract location if pattern like "City, ST"
                    loc_match = re.search(r'[-–—]\s*([A-Z][a-z]+,\s*[A-Z]{2}(?:,\s*[A-Z]+)?)', line)
                    if loc_match:
                        edu.location = loc_match.group(1).strip()

                    if edu.institution:
                        education_list.append(edu)

            i += 1  # Move to next line

        return education_list

    def _split_skills_preserve_groups(self, text: str, delimiter: str) -> List[str]:
        """Split skills by delimiter but preserve content within parentheses as single units."""
        skills = []
        current = ""
        paren_depth = 0

        for char in text:
            if char == '(':
                paren_depth += 1
                current += char
            elif char == ')':
                paren_depth -= 1
                current += char
            elif char == delimiter and paren_depth == 0:
                # Split here
                skill = current.strip()
                if skill:
                    skills.append(skill)
                current = ""
            else:
                current += char

        # Add the last skill
        skill = current.strip()
        if skill:
            skills.append(skill)

        return skills

    def _parse_skills(self, text: str) -> List[str]:
        """Parse skills from text with timeout protection."""
        if not text.strip():
            return []

        # Use simple parser for now (detailed parser has regex issues causing hangs)
        print(f"[PARSER] Using simple skills parser")
        return self._parse_skills_simple(text)

    def _parse_skills_simple(self, text: str) -> List[str]:
        """Simple fallback skills parser that extracts individual skills."""
        skills = []
        # Remove bullet points and asterisks
        text = re.sub(r'^[•\-*▪●○]\s*', '', text, flags=re.MULTILINE)
        text = re.sub(r'\*\*', '', text)  # Remove markdown bold

        # Common skill category prefixes to remove
        category_prefixes = [
            'cloud platforms?', 'programming languages?', 'databases?', 'big data',
            'processing', 'analytics', 'devops', 'ci/cd', 'tools?', 'frameworks?',
            'genai', 'responsible ai', 'methods?', 'other tools?', 'orchestration',
            'data', 'skills', 'technologies', 'platforms?'
        ]

        # Create regex pattern for category prefixes
        category_pattern = r'^(' + '|'.join(category_prefixes) + r')[\s&]+(.+)$'

        # Split by lines
        for line in text.split('\n'):
            line = line.strip()
            if not line or len(line) < 2:
                continue

            # If line has category with colon (e.g., "Cloud Platforms: Azure, AWS")
            if ':' in line:
                # Take only the skills part (after colon)
                skills_part = line.split(':', 1)[1].strip()
            else:
                # Check if line starts with a category prefix (without colon)
                # e.g., "Cloud Platforms Azure Data Factory"
                category_match = re.match(category_pattern, line, re.IGNORECASE)
                if category_match:
                    # Extract just the skills part (after category)
                    skills_part = category_match.group(2).strip()
                else:
                    skills_part = line

            # Split by common delimiters (comma, pipe, semicolon)
            for skill in re.split(r'[,|;]', skills_part):
                skill = skill.strip()
                # Clean up parentheses and extra text
                skill = re.sub(r'\s*\([^)]*\)', '', skill)  # Remove (ADF) etc.
                skill = skill.strip()

                # Filter out noise and keep only meaningful skills
                if skill and len(skill) >= 2 and len(skill) <= 50:
                    # Skip common noise words
                    noise_words = ['skills', 'technologies', 'tools', 'platforms', 'others', 'frameworks']
                    if skill.lower() not in noise_words:
                        skills.append(skill)

        # Remove duplicates while preserving order
        seen = set()
        unique_skills = []
        for skill in skills:
            skill_lower = skill.lower()
            if skill_lower not in seen:
                seen.add(skill_lower)
                unique_skills.append(skill)

        print(f"[PARSER] Extracted {len(unique_skills)} individual skills")
        for i, skill in enumerate(unique_skills[:10], 1):
            print(f"[PARSER]   {i}. {skill}")
        if len(unique_skills) > 10:
            print(f"[PARSER]   ... and {len(unique_skills) - 10} more")

        return unique_skills[:100]  # Limit to 100 skills

    def _parse_skills_detailed(self, text: str) -> List[str]:
        """Detailed skills parser (original implementation)."""
        skills = []

        # Remove bullet points and clean up
        text = re.sub(r'^[•\-*▪●○]\s*', '', text, flags=re.MULTILINE)

        # Remove common noise phrases from PDF extraction
        noise_phrases = [
            r'Page \d+ of \d+',
            r'Some skills have been grouped together.*?(?=\n|$)',
            r'GEN AI SKILL SET.*?(?=\n|$)',
            r'TECHNICAL SKILLS.*?(?=\n|$)',
            r'SKILLS:.*?(?=\n|$)',
            r'\(Learning\)',
            r'Proficiency',
            r'Applications',
            r'Skill Category',
            r'Skills & Technologies',
            r'Beginner(?:\s|,|$)',
        ]
        for phrase in noise_phrases:
            text = re.sub(phrase, '', text, flags=re.IGNORECASE)

        # Split by lines first
        lines = text.split('\n')

        for line in lines:
            line = line.strip()
            if not line or len(line) < 3:
                continue

            # Check if line has category pattern "Category: skills"
            if ':' in line and not line.startswith('http'):
                # Handle multiple categories on same line
                # Category keywords that indicate a new category section
                category_keywords = [
                    'Languages', 'Tools', 'Platforms', 'Technologies', 'Frameworks',
                    'Management', 'Warehousing', 'Integration', 'Visualization',
                    'Code', 'Methodologies', 'Skills'
                ]

                # Build pattern that matches these keywords at end of category name
                # Pattern: "Some Words (Language|Tool|Platform|...): "
                keyword_pattern = '|'.join(category_keywords)
                category_pattern = rf'([A-Za-z\s]+(?:{keyword_pattern})):\s*'
                segments = re.split(category_pattern, line)

                # If we got meaningful segments, process them
                if len(segments) > 2:
                    for i in range(1, len(segments), 2):
                        if i + 1 < len(segments):
                            skills_part = segments[i + 1].strip()

                            # Split skills by comma, BUT preserve parentheses groups
                            if ',' in skills_part:
                                line_skills = self._split_skills_preserve_groups(skills_part, ',')
                            elif '|' in skills_part:
                                line_skills = self._split_skills_preserve_groups(skills_part, '|')
                            else:
                                line_skills = [skills_part] if skills_part else []

                            skills.extend(line_skills)
                else:
                    # Simple case - single category: skills pattern
                    parts = line.split(':', 1)
                    if len(parts) == 2:
                        skills_part = parts[1].strip()

                        # Split skills by comma, BUT preserve parentheses groups
                        if ',' in skills_part:
                            line_skills = self._split_skills_preserve_groups(skills_part, ',')
                        elif '|' in skills_part:
                            line_skills = self._split_skills_preserve_groups(skills_part, '|')
                        else:
                            line_skills = [skills_part] if skills_part else []

                        skills.extend(line_skills)
            else:
                # No category, try to split by delimiters
                if ',' in line:
                    line_skills = self._split_skills_preserve_groups(line, ',')
                elif '|' in line:
                    line_skills = self._split_skills_preserve_groups(line, '|')
                else:
                    # Single skill on the line
                    line_skills = [line]

                skills.extend(line_skills)

        # Clean up and filter
        cleaned_skills = []
        for skill in skills:
            # Remove any remaining newlines or excess whitespace
            skill = ' '.join(skill.split())

            # Remove parenthetical notes like "(exposure)" or "(PySpark)"
            # But keep them if they're part of the skill name
            skill = re.sub(r'\s*\([^)]*exposure[^)]*\)', '', skill, flags=re.IGNORECASE)

            # Filter out empty, too short, or too long entries
            if skill and 2 < len(skill) < 80:
                # Skip section headers and noise
                skip_terms = [
                    'skills:', 'proficiency', 'applications', 'skill category',
                    'technical skills', 'gen ai skill', 'page', 'learning',
                    'custom assistants', 'enterprise solution', 'semantic search',
                    'similarity matching', 'vector store', 'document q&a',
                    'knowledge retrieval', 'intelligent', 'specialized domain',
                    'consistent llm', 'reasoning tasks', 'production-grade',
                    'compliance with', 'enterprise-grade', 'some skills have'
                ]

                skill_lower = skill.lower()
                if not any(term in skill_lower for term in skip_terms):
                    # Only add if it has substantial content
                    if len(skill.split()) <= 10:  # Max 10 words for a skill
                        cleaned_skills.append(skill)

        # Remove duplicates while preserving order
        seen = set()
        unique_skills = []
        for skill in cleaned_skills:
            skill_lower = skill.lower()
            if skill_lower not in seen and skill_lower not in ['beginner', 'is', 'and', 'or', 'the']:
                seen.add(skill_lower)
                unique_skills.append(skill)

        return unique_skills

    def _parse_genai_skills(self, text: str) -> List[str]:
        """Parse GenAI/ML skills from GEN AI SKILL SET section."""
        if not text.strip():
            return []

        # GenAI skills are often in a table format with columns
        # We want to extract actual skill names and technologies
        genai_skills = []

        # Remove common table headers and noise
        noise_patterns = [
            r'Skill Category',
            r'Skills & Technologies',
            r'Proficiency',
            r'Applications',
            r'Beginner',
            r'Intermediate',
            r'Advanced',
            r'\(Learning\)',
            r'GEN AI SKILL SET.*?(?=\n|$)',
        ]
        for pattern in noise_patterns:
            text = re.sub(pattern, '', text, flags=re.IGNORECASE)

        # Known GenAI skill keywords to extract
        genai_keywords = [
            # Foundation Models
            'LLM', 'LLMs', 'GPT', 'Claude', 'Llama', 'Mistral', 'Midjourney', 'DALL-E',
            'fine-tuning', 'model training', 'foundation models',
            # Vector Databases
            'Pinecone', 'Chroma', 'FAISS', 'Weaviate', 'Qdrant', 'Milvus',
            'Azure AI Search', 'vector database', 'vector store',
            # RAG
            'RAG', 'Retrieval Augmented Generation', 'chunking', 'embedding',
            'context window', 'hybrid search', 'semantic search',
            # Prompt Engineering
            'prompt engineering', 'CoT', 'chain of thought', 'few-shot learning',
            'zero-shot', 'system prompt', 'prompt design', 'output control',
            # Architectures
            'Transformers', 'attention mechanism', 'BERT', 'T5',
            'embedding models', 'token processing', 'tokenization',
            # LLM Frameworks
            'LangChain', 'LlamaIndex', 'Haystack', 'Azure OpenAI',
            'OpenAI API', 'Anthropic', 'AWS Bedrock', 'Vertex AI',
            'MLflow', 'Weights & Biases', 'Hugging Face',
            # Responsible AI
            'responsible AI', 'AI ethics', 'hallucination detection',
            'output safety', 'privacy techniques', 'bias mitigation',
            'model governance', 'AI compliance',
        ]

        # Extract skills using the general skill parsing logic first
        parsed_skills = self._parse_skills(text)

        # Filter and keep only GenAI-relevant skills
        for skill in parsed_skills:
            skill_lower = skill.lower()

            # Check if skill matches any GenAI keyword
            is_genai = any(keyword.lower() in skill_lower for keyword in genai_keywords)

            # Or if skill contains key GenAI terms
            genai_terms = ['llm', 'genai', 'ai', 'ml', 'rag', 'vector', 'embedding',
                          'prompt', 'transformer', 'langchain', 'openai', 'azure openai']
            contains_genai_term = any(term in skill_lower for term in genai_terms)

            if is_genai or contains_genai_term:
                # Skip very generic or noise terms
                skip_terms = ['similarity', 'matching', 'integration', 'knowledge',
                             'retrieval', 'selection', 'optimization', 'search',
                             'design', 'output', 'control', 'adaptation', 'assistants',
                             'reasoning', 'tasks', 'architecture', 'custom', 'mechanisms',
                             'performance', 'models', 'processing', 'detection',
                             'safety', 'techniques', 'policies', 'frameworks', 'systems']

                if skill_lower not in skip_terms and len(skill) > 2:
                    genai_skills.append(skill)

        # Add common GenAI skills that might have been in the text
        for keyword in genai_keywords:
            if keyword.lower() in text.lower() and keyword not in genai_skills:
                # Check if it's actually present as a distinct skill
                if re.search(r'\b' + re.escape(keyword) + r'\b', text, re.IGNORECASE):
                    genai_skills.append(keyword)

        # Remove duplicates while preserving order
        seen = set()
        unique_genai_skills = []
        for skill in genai_skills:
            skill_lower = skill.lower()
            if skill_lower not in seen:
                seen.add(skill_lower)
                unique_genai_skills.append(skill)

        return unique_genai_skills[:30]  # Limit to top 30 to avoid noise

    def _parse_projects(self, text: str) -> List[Project]:
        """Parse projects from text."""
        projects = []

        if not text.strip():
            return projects

        # First, find where "PROJECT HIGHLIGHTS" appears and only parse from there
        project_header_match = re.search(r'PROJECT\s+HIGHLIGHTS?', text, flags=re.IGNORECASE)
        if project_header_match:
            # Only parse text from the PROJECT HIGHLIGHTS header onwards
            text = text[project_header_match.end():].strip()
        else:
            # If no header found, we might be in the wrong section
            # Check if text starts with bullets (likely misidentified experience bullets)
            lines = text.split('\n')
            first_non_empty = next((line for line in lines if line.strip()), "")
            if first_non_empty.startswith(('•', '-', '*', '▪', '●', '○')):
                # This looks like experience bullets, not projects
                # Skip lines until we find a project-like title
                found_project_start = False
                start_idx = 0
                for i, line in enumerate(lines):
                    line = line.strip()
                    # Look for a line that doesn't start with bullet and has project-like format
                    if line and not line.startswith(('•', '-', '*', '▪', '●', '○')) and len(line) > 10 and '–' in line:
                        start_idx = i
                        found_project_start = True
                        break

                if found_project_start:
                    text = '\n'.join(lines[start_idx:])
                else:
                    # No projects found, return empty
                    return projects

        lines = text.split('\n')
        current_project = None

        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue

            # Skip page markers
            if re.match(r'Page\s+\d+\s+of\s+\d+', line, re.IGNORECASE):
                continue

            # Check if this is a project title (not a bullet point)
            # Project titles usually don't start with bullet markers and have substantial length
            if not line.startswith(('•', '-', '*', '▪', '●', '○')):
                # Skip section headers
                if 'project' in line.lower() and 'highlight' in line.lower():
                    continue

                # This might be a project title
                # Check if it's not too short and looks like a title
                # Must either have "–" (subtitle) OR be followed by a bullet point
                if len(line) > 15 and (
                    '–' in line or  # Has subtitle separator like "Project Name – Subtitle"
                    (i + 1 < len(lines) and lines[i + 1].strip().startswith(('•', '-', '*', '●')))  # Next line is a bullet
                ):
                    # Save previous project
                    if current_project and current_project.bullets:
                        projects.append(current_project)

                    # Start new project
                    current_project = Project(name=line)
                    current_project.bullets = []

            else:
                # This is a bullet point
                if current_project:
                    cleaned = re.sub(r'^[•\-*▪●○]\s*', '', line)
                    if cleaned and len(cleaned) > 5:
                        current_project.bullets.append(cleaned)

        # Add last project
        if current_project and current_project.bullets:
            projects.append(current_project)

        return projects

    def _parse_certifications(self, text: str) -> List[Certification]:
        """Parse certifications from text."""
        certifications = []

        if not text.strip():
            return certifications

        lines = [line.strip() for line in text.split('\n') if line.strip()]

        for line in lines:
            # Remove bullet points
            line = re.sub(r'^[•\-*▪]\s*', '', line)

            # Skip empty or very short lines
            if len(line) < 3:
                continue

            # Try to extract issuer and date
            # Pattern: "Certification Name - Issuer (Year)" or similar
            # BUT preserve hyphens in cert codes like "AZ-900"
            # Split only on spaced hyphens or commas
            parts = re.split(r'\s+[-–—]\s+|,', line)

            cert_name = parts[0].strip() if parts else line
            cert_issuer = parts[1].strip() if len(parts) > 1 else ""

            cert = Certification(
                name=cert_name,
                issuer=cert_issuer
            )

            # Look for year
            year_match = re.search(r'\b\d{4}\b', line)
            if year_match:
                cert.date = year_match.group(0)

            certifications.append(cert)

        return certifications
