"""DOCX formatter for exporting resumes and cover letters."""

from typing import List, Optional, Dict
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from src.models import Resume, Experience, Education
import re


class DOCXFormatter:
    """Format and export documents to DOCX with ATS-friendly styling."""

    def __init__(self, template_style: str = "professional"):
        """
        Initialize DOCX formatter.

        Args:
            template_style: Template style (professional, modern, minimal)
        """
        self.template_style = template_style
        self.doc = None  # Will be created per export

    def _create_new_document(self):
        """Create a new document with proper styling."""
        self.doc = Document()
        self._setup_document_styles()

    def _setup_document_styles(self):
        """Set up document-wide styles."""
        # Set normal style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Set margins (ATS-friendly: 0.5-1 inch)
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

    def _add_heading(self, text: str, level: int = 1, color: str = "000000"):
        """
        Add a formatted heading.

        Args:
            text: Heading text
            level: Heading level (1=name, 2=section, 3=subsection)
            color: Hex color code
        """
        if level == 1:  # Name heading
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(text.upper())
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(color)
            para.space_after = Pt(6)

        elif level == 2:  # Section heading
            para = self.doc.add_paragraph()
            run = para.add_run(text.upper())
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(color)

            # Add underline
            self._add_paragraph_border(para, border_color=color)
            para.space_after = Pt(6)

        elif level == 3:  # Subsection (company, role)
            para = self.doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(11)
            run.font.bold = True
            para.space_after = Pt(3)

    def _add_paragraph_border(self, paragraph, border_color: str = "000000"):
        """Add bottom border to paragraph."""
        p = paragraph._element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')

        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')  # Border size
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), border_color)

        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_hyperlink(self, paragraph, text: str, url: str):
        """
        Add a hyperlink to a paragraph.

        Args:
            paragraph: Paragraph object
            text: Display text
            url: URL to link to
        """
        # Ensure URL has protocol
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url

        # Create hyperlink element
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        # Create run element
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Style as hyperlink (blue, underlined)
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')
        rPr.append(rStyle)
        new_run.append(rPr)

        # Add text
        new_run.text = text
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)

        return hyperlink

    def _add_contact_info(self, resume: Resume):
        """Add contact information section with clickable hyperlinks."""
        contact = resume.contact

        # Center-aligned contact info
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Build contact info with hyperlinks
        first = True

        if contact.email:
            if not first:
                run = para.add_run(" | ")
                run.font.size = Pt(10)
            self._add_hyperlink(para, contact.email, f"mailto:{contact.email}")
            first = False

        if contact.phone:
            if not first:
                run = para.add_run(" | ")
                run.font.size = Pt(10)
            run = para.add_run(contact.phone)
            run.font.size = Pt(10)
            first = False

        if contact.location:
            if not first:
                run = para.add_run(" | ")
                run.font.size = Pt(10)
            run = para.add_run(contact.location)
            run.font.size = Pt(10)
            first = False

        if contact.linkedin:
            if not first:
                run = para.add_run(" | ")
                run.font.size = Pt(10)
            self._add_hyperlink(para, "LinkedIn", contact.linkedin)
            first = False

        if contact.github:
            if not first:
                run = para.add_run(" | ")
                run.font.size = Pt(10)
            self._add_hyperlink(para, "GitHub", contact.github)
            first = False

        if contact.portfolio:
            if not first:
                run = para.add_run(" | ")
                run.font.size = Pt(10)
            self._add_hyperlink(para, "Portfolio", contact.portfolio)
            first = False

        if contact.website:
            if not first:
                run = para.add_run(" | ")
                run.font.size = Pt(10)
            self._add_hyperlink(para, "Website", contact.website)
            first = False

        para.space_after = Pt(12)

    def _add_bullet_point(self, text: str, bold_keywords: List[str] = None):
        """
        Add a bullet point with optional keyword bolding or markdown bold parsing.

        Args:
            text: Bullet text (may contain **markdown** bold syntax)
            bold_keywords: Keywords to bold
        """
        para = self.doc.add_paragraph(style='List Bullet')
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.space_after = Pt(3)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justify bullet text

        # Always check for markdown bold first
        if '**' in text:
            self._add_text_with_markdown_bold(para, text)
        elif bold_keywords:
            # Split text and bold keywords
            self._add_text_with_bold_keywords(para, text, bold_keywords)
        else:
            run = para.add_run(text)
            run.font.size = Pt(11)

    def _add_text_with_markdown_bold(self, paragraph, text: str, font_size: int = 11):
        """
        Add text with markdown bold syntax (**text**) parsed and applied.

        Args:
            paragraph: Paragraph object
            text: Text with markdown bold syntax
            font_size: Font size in points
        """
        # Pattern to match **text**
        pattern = r'\*\*(.+?)\*\*'

        # Split text by bold markers
        parts = re.split(pattern, text)

        is_bold = False
        for i, part in enumerate(parts):
            if part:
                run = paragraph.add_run(part)
                run.font.size = Pt(font_size)

                # Alternate between regular and bold
                # Even indices are regular text, odd indices are bold
                if i % 2 == 1:  # Odd index = captured group = bold text
                    run.font.bold = True

    def _add_text_with_bold_keywords(self, paragraph, text: str, keywords: List[str]):
        """Add text with specific keywords bolded."""
        # First check if text has markdown bold syntax
        if '**' in text:
            self._add_text_with_markdown_bold(paragraph, text)
            return

        # Otherwise, use keyword-based bolding
        if not keywords:
            run = paragraph.add_run(text)
            run.font.size = Pt(11)
            return

        # Sort keywords by length (longest first) to avoid partial matches
        sorted_keywords = sorted(keywords, key=len, reverse=True)

        # Escape special regex characters
        escaped_keywords = [re.escape(kw) for kw in sorted_keywords]

        # Create pattern
        pattern = '|'.join(escaped_keywords)

        if not pattern:
            run = paragraph.add_run(text)
            run.font.size = Pt(11)
            return

        # Split text by keywords
        parts = re.split(f'({pattern})', text, flags=re.IGNORECASE)

        for part in parts:
            if part:
                run = paragraph.add_run(part)
                run.font.size = Pt(11)

                # Bold if it's a keyword
                if any(part.lower() == kw.lower() for kw in sorted_keywords):
                    run.font.bold = True

    def _parse_skills_categories(self, skills: List[str]) -> Dict[str, str]:
        """
        Parse skills list into categories.

        Args:
            skills: List of skill lines (may include category headers with **)

        Returns:
            Dictionary mapping category names to comma-separated skills
        """
        categories = {}
        current_category = None

        for line in skills:
            line = line.strip()
            if not line:
                continue

            # Check if this is a category header (contains ** or ends with :)
            if '**' in line:
                # Extract category name (remove **)
                category = line.replace('**', '').strip()
                # Remove trailing colon if present
                if ':' in category:
                    parts = category.split(':', 1)
                    current_category = parts[0].strip()
                    # If there are skills after the colon, add them
                    if len(parts) > 1 and parts[1].strip():
                        categories[current_category] = parts[1].strip()
                    else:
                        categories[current_category] = ""
                else:
                    current_category = category
                    categories[current_category] = ""
            elif current_category:
                # This line contains skills for the current category
                if categories[current_category]:
                    categories[current_category] += ", " + line
                else:
                    categories[current_category] = line
            else:
                # No category yet, create "Other" category
                if "Other" not in categories:
                    categories["Other"] = line
                else:
                    categories["Other"] += ", " + line

        # Clean up empty categories
        categories = {k: v for k, v in categories.items() if v}

        return categories

    def _format_date_range(self, start_date: str, end_date: str, is_current: bool) -> str:
        """Format date range."""
        if is_current:
            return f"{start_date} – Present"
        return f"{start_date} – {end_date}"

    def export_resume(
        self,
        resume: Resume,
        output_path: str,
        bold_keywords: List[str] = None
    ):
        """
        Export resume to DOCX.

        Args:
            resume: Resume object
            output_path: Output file path
            bold_keywords: Keywords to bold throughout document
        """
        print(f"\nExporting resume to DOCX: {output_path}")

        # Create a fresh document for this export
        self._create_new_document()

        # Add name
        self._add_heading(resume.contact.full_name, level=1, color="1F4788")

        # Add contact info
        self._add_contact_info(resume)

        # Professional Summary
        if resume.summary:
            self._add_heading("Professional Summary", level=2, color="1F4788")
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justify text

            # Parse markdown bold in summary
            if '**' in resume.summary:
                self._add_text_with_markdown_bold(para, resume.summary)
            else:
                run = para.add_run(resume.summary)
                run.font.size = Pt(11)

            para.space_after = Pt(12)

        # Technical Skills
        if resume.skills:
            self._add_heading("Technical Skills", level=2, color="1F4788")

            # Check if skills are in categorized format (has ** headers)
            # Parse skills into categories
            if resume.skills and '**' in str(resume.skills[0]):
                # Categorized format - create a table
                categories = self._parse_skills_categories(resume.skills)

                if categories:
                    # Create table with 2 columns: Category | Skills
                    table = self.doc.add_table(rows=len(categories), cols=2)
                    table.style = 'Table Grid'

                    # Set column widths
                    table.autofit = False
                    table.allow_autofit = False
                    widths = (Inches(1.5), Inches(4.5))

                    for row_idx, (category, skills_list) in enumerate(categories.items()):
                        row = table.rows[row_idx]

                        # Category cell (left column)
                        category_cell = row.cells[0]
                        category_para = category_cell.paragraphs[0]
                        category_para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Align category left
                        category_run = category_para.add_run(category)
                        category_run.font.bold = True
                        category_run.font.size = Pt(11)
                        category_cell.width = widths[0]

                        # Skills cell (right column)
                        skills_cell = row.cells[1]
                        skills_para = skills_cell.paragraphs[0]
                        skills_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justify skills text
                        skills_run = skills_para.add_run(skills_list)
                        skills_run.font.size = Pt(11)
                        skills_cell.width = widths[1]

                    # Add space after table
                    self.doc.add_paragraph().space_after = Pt(12)
                else:
                    # Fallback to paragraph format if parsing fails
                    for skill_line in resume.skills:
                        para = self.doc.add_paragraph()
                        if '**' in skill_line:
                            self._add_text_with_markdown_bold(para, skill_line)
                        else:
                            run = para.add_run(skill_line)
                            run.font.size = Pt(11)
                        para.space_after = Pt(3)
            else:
                # Simple comma-separated list
                skills_text = ", ".join(resume.skills)
                para = self.doc.add_paragraph()

                if '**' in skills_text:
                    self._add_text_with_markdown_bold(para, skills_text)
                elif bold_keywords:
                    self._add_text_with_bold_keywords(para, skills_text, bold_keywords)
                else:
                    run = para.add_run(skills_text)
                    run.font.size = Pt(11)

                para.space_after = Pt(12)

        # GenAI/ML Skills (Separate Section)
        if resume.genai_skills:
            self._add_heading("GenAI & Machine Learning Skills", level=2, color="1F4788")

            # Check if GenAI skills are in categorized format
            if resume.genai_skills and '**' in str(resume.genai_skills[0]):
                # Categorized format - create a table
                genai_categories = self._parse_skills_categories(resume.genai_skills)

                if genai_categories:
                    # Create table with 2 columns: Category | Skills
                    table = self.doc.add_table(rows=len(genai_categories), cols=2)
                    table.style = 'Table Grid'

                    # Set column widths
                    table.autofit = False
                    table.allow_autofit = False
                    widths = (Inches(1.5), Inches(4.5))

                    for row_idx, (category, skills_list) in enumerate(genai_categories.items()):
                        row = table.rows[row_idx]

                        # Category cell (left column)
                        category_cell = row.cells[0]
                        category_para = category_cell.paragraphs[0]
                        category_para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Align category left
                        category_run = category_para.add_run(category)
                        category_run.font.bold = True
                        category_run.font.size = Pt(11)
                        category_cell.width = widths[0]

                        # Skills cell (right column)
                        skills_cell = row.cells[1]
                        skills_para = skills_cell.paragraphs[0]
                        skills_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justify skills text
                        skills_run = skills_para.add_run(skills_list)
                        skills_run.font.size = Pt(11)
                        skills_cell.width = widths[1]

                    # Add space after table
                    self.doc.add_paragraph().space_after = Pt(12)
                else:
                    # Fallback to paragraph format
                    for skill_line in resume.genai_skills:
                        para = self.doc.add_paragraph()
                        if '**' in skill_line:
                            self._add_text_with_markdown_bold(para, skill_line)
                        else:
                            run = para.add_run(skill_line)
                            run.font.size = Pt(11)
                        para.space_after = Pt(3)
            else:
                # Simple comma-separated list
                genai_skills_text = ", ".join(resume.genai_skills)
                para = self.doc.add_paragraph()

                if '**' in genai_skills_text:
                    self._add_text_with_markdown_bold(para, genai_skills_text)
                elif bold_keywords:
                    self._add_text_with_bold_keywords(para, genai_skills_text, bold_keywords)
                else:
                    run = para.add_run(genai_skills_text)
                    run.font.size = Pt(11)

                para.space_after = Pt(12)

        # Professional Experience
        if resume.experience:
            self._add_heading("Professional Experience", level=2, color="1F4788")

            for exp in resume.experience:
                # Title - only add if not "Unknown"
                if exp.title and exp.title.lower() != "unknown":
                    self._add_heading(f"{exp.title}", level=3)

                # Company, location, dates
                para = self.doc.add_paragraph()
                company_text = exp.company
                if exp.location:
                    company_text += f", {exp.location}"
                company_text += f" | {self._format_date_range(exp.start_date, exp.end_date, exp.is_current)}"

                run = para.add_run(company_text)
                run.font.size = Pt(10)
                run.font.italic = True
                para.space_after = Pt(6)

                # Bullets
                for bullet in exp.bullets:
                    self._add_bullet_point(bullet, bold_keywords)

                # Add space after experience
                self.doc.add_paragraph().space_after = Pt(6)

        # Education
        if resume.education:
            self._add_heading("Education", level=2, color="1F4788")

            for edu in resume.education:
                # First line: Degree in Field of Study (if available)
                para = self.doc.add_paragraph()

                degree_parts = []
                if edu.degree:
                    degree_parts.append(edu.degree)
                if edu.field_of_study:
                    degree_parts.append(f"in {edu.field_of_study}")

                if degree_parts:
                    degree_text = " ".join(degree_parts)
                    run = para.add_run(degree_text)
                    run.font.size = Pt(11)
                    run.font.bold = True

                # Second line: Institution, Location (if available)
                para_inst = self.doc.add_paragraph()
                inst_parts = []
                if edu.institution:
                    inst_parts.append(edu.institution)
                if edu.location:
                    inst_parts.append(edu.location)

                if inst_parts:
                    inst_text = ", ".join(inst_parts)
                    run_inst = para_inst.add_run(inst_text)
                    run_inst.font.size = Pt(10)
                    run_inst.font.italic = True

                # Third line: Graduation Date | GPA (if available)
                details_parts = []
                if edu.graduation_date:
                    details_parts.append(f"Graduated: {edu.graduation_date}")
                if edu.gpa:
                    details_parts.append(f"GPA: {edu.gpa}")

                if details_parts:
                    para_details = self.doc.add_paragraph()
                    details_text = " | ".join(details_parts)
                    run_details = para_details.add_run(details_text)
                    run_details.font.size = Pt(10)
                    run_details.font.color.rgb = RGBColor(0x4A, 0x5C, 0x6A)  # Subtle gray

                para.space_after = Pt(6)

        # Certifications
        if resume.certifications:
            self._add_heading("Certifications", level=2, color="1F4788")

            for cert in resume.certifications:
                para = self.doc.add_paragraph(style='List Bullet')

                # Handle both string and Certification object
                if isinstance(cert, str):
                    cert_text = cert
                else:
                    cert_text = cert.name if hasattr(cert, 'name') else str(cert)
                    if hasattr(cert, 'issuer') and cert.issuer:
                        cert_text += f" - {cert.issuer}"
                    if hasattr(cert, 'date') and cert.date:
                        cert_text += f" ({cert.date})"

                run = para.add_run(cert_text)
                run.font.size = Pt(11)

        # Save document
        self.doc.save(output_path)
        print(f"✓ Resume exported successfully")

    def export_cover_letter(
        self,
        cover_letter_text: str,
        candidate_name: str,
        output_path: str
    ):
        """
        Export cover letter to DOCX.

        Args:
            cover_letter_text: Cover letter text
            candidate_name: Candidate's name
            output_path: Output file path
        """
        print(f"\nExporting cover letter to DOCX: {output_path}")

        # Create a fresh document for this export
        self._create_new_document()

        # Add name at top
        para_name = self.doc.add_paragraph()
        para_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_name = para_name.add_run(candidate_name)
        run_name.font.size = Pt(14)
        run_name.font.bold = True
        para_name.space_after = Pt(12)

        # Add current date
        from datetime import datetime
        para_date = self.doc.add_paragraph()
        run_date = para_date.add_run(datetime.now().strftime("%B %d, %Y"))
        run_date.font.size = Pt(11)
        para_date.space_after = Pt(24)

        # Add cover letter content
        # Split by paragraphs
        paragraphs = cover_letter_text.split('\n\n')

        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                para = self.doc.add_paragraph(para_text)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para_format = para.paragraph_format
                para_format.space_after = Pt(12)
                para_format.line_spacing = 1.15

                for run in para.runs:
                    run.font.size = Pt(11)

        # Save document
        self.doc.save(output_path)
        print(f"✓ Cover letter exported successfully")

    def export_combined_document(
        self,
        resume: Resume,
        cover_letter_text: str,
        output_path: str,
        bold_keywords: List[str] = None
    ):
        """
        Export cover letter and resume in one document.

        Args:
            resume: Resume object
            cover_letter_text: Cover letter text
            output_path: Output file path
            bold_keywords: Keywords to bold in resume
        """
        print(f"\nExporting combined document to DOCX: {output_path}")

        # First export cover letter
        self.export_cover_letter(cover_letter_text, resume.contact.full_name, output_path)

        # Add page break
        self.doc.add_page_break()

        # Then export resume on second page
        # Add name
        self._add_heading(resume.contact.full_name, level=1, color="1F4788")
        self._add_contact_info(resume)

        # Continue with resume sections...
        # (Same as export_resume but without recreating document)

        self.doc.save(output_path)
        print(f"✓ Combined document exported successfully")


class DOCXExporter:
    """High-level DOCX export interface."""

    @staticmethod
    def export_resume(
        resume: Resume,
        output_path: str,
        template_style: str = "professional",
        bold_keywords: List[str] = None
    ):
        """
        Export resume to DOCX file.

        Args:
            resume: Resume object
            output_path: Output file path
            template_style: Template style
            bold_keywords: Keywords to bold
        """
        formatter = DOCXFormatter(template_style)
        formatter.export_resume(resume, output_path, bold_keywords)

    @staticmethod
    def export_cover_letter(
        cover_letter_text: str,
        candidate_name: str,
        output_path: str,
        template_style: str = "professional"
    ):
        """
        Export cover letter to DOCX file.

        Args:
            cover_letter_text: Cover letter text
            candidate_name: Candidate's name
            output_path: Output file path
            template_style: Template style
        """
        formatter = DOCXFormatter(template_style)
        formatter.export_cover_letter(cover_letter_text, candidate_name, output_path)

    @staticmethod
    def export_combined(
        resume: Resume,
        cover_letter_text: str,
        output_path: str,
        template_style: str = "professional",
        bold_keywords: List[str] = None
    ):
        """
        Export cover letter + resume in one document.

        Args:
            resume: Resume object
            cover_letter_text: Cover letter text
            output_path: Output file path
            template_style: Template style
            bold_keywords: Keywords to bold
        """
        formatter = DOCXFormatter(template_style)
        formatter.export_combined_document(resume, cover_letter_text, output_path, bold_keywords)
